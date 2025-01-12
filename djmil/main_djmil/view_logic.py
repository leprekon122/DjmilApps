"""import pockets"""
import csv
from datetime import datetime, timedelta
from decimal import Decimal
from docx import Document
from django.http import HttpResponse
from django.db.models import Count
from django.db.models import Q

from .models import SecondOrdersModel, MainOrders, FlightRecorderModel, SkySafeData


class SecondOnlineSQLReq:
    """SecondOnlineSQLReq class"""

    def __init__(self, *args):
        self.drone_id = args

    @property
    def make_sql(self):
        '''make sql query'''
        return SecondOrdersModel.objects.all()

    @property
    def search_by_drone_id(self):
        """logic for filter in second_online orders page"""
        return SecondOrdersModel.objects.filter(serial_no=self.drone_id[0]).order_by('-dt')


class OnlineSQLReq:
    """Logic for Online sql req"""

    def __init__(self, *args):
        self.drone_id = args

    @property
    def standart_req(self):
        """function for first req in OnlineOrder"""
        return MainOrders.objects.all()

    @property
    def newest_req(self):
        """filter logic for newest data"""
        return MainOrders.objects.all().order_by('-dt_last')

    @property
    def oldest_req(self):
        """filter logic for oldest data"""
        return MainOrders.objects.all().order_by('dt_first')

    @property
    def search_drone_id(self):
        """searching by drone_id"""
        return MainOrders.objects.filter(serial_no=self.drone_id[0])


class DownloadOnlineOrders:
    """class for downloading order """

    @property
    def download_order(self):
        """logic for download order in Online Second Order page """
        model = MainOrders.objects.values()
        response = HttpResponse(
            content_type='text/csv',
            headers={'Content-Disposition': 'attachment; filename="somefilename.csv"'},
        )

        writer = csv.writer(response)
        writer.writerow(['id', 'serial_no', 'product_type', 'dt_first', 'dt_last'])
        for elem in model:
            writer.writerow(elem.values())

        return response

    @property
    def download_newest_order(self):
        """function for downloading newest order """
        model = MainOrders.objects.values().order_by('-dt_last')
        response = HttpResponse(
            content_type='text/csv',
            headers={'Content-Disposition': 'attachment; filename="somefilename.csv"'},
        )

        writer = csv.writer(response)
        writer.writerow(['id', 'serial_no', 'product_type', 'dt_first', 'dt_last'])
        for el in model:
            writer.writerow(el.values())

        return response

    @property
    def download_oldest_order(self):
        """function for downloading oldest order """
        model = MainOrders.objects.values().order_by('id')
        response = HttpResponse(
            content_type='text/csv',
            headers={'Content-Disposition': 'attachment; filename="somefilename.csv"'},
        )

        writer = csv.writer(response)
        writer.writerow(['id', 'serial_no', 'product_type', 'dt_first', 'dt_last'])
        for elem in model:
            writer.writerow(elem.values())

        return response


class DownloadSecondOnlineOrders:
    """download second orders on production """

    @property
    def download_order(self):
        """function for downloading order"""
        model = SecondOrdersModel.objects.values()
        response = HttpResponse(
            content_type='text/csv',
            headers={'Content-Disposition': 'attachment; filename="somefilename.csv"'},
        )

        writer = csv.writer(response)
        writer.writerow(['SELECT serial_no, product_type, longitude, latitude, height,'
                         ' altitude, phone_app_latitude, phone_app_longitude, phone_app_x,'
                         ' phone_app_y, home_latitude,'
                         ' home_longitude, home_x, home_y, dt, frame_id, status'])
        for elem in model:
            writer.writerow(elem.values())

        return response

    @property
    def download_newest_order(self):
        """function for downloading by newest filter"""
        model = SecondOrdersModel.objects.values().order_by('-dt')
        response = HttpResponse(
            content_type='text/csv',
            headers={'Content-Disposition': 'attachment; filename="somefilename.csv"'},
        )

        writer = csv.writer(response)
        writer.writerow(['SELECT serial_no, product_type, longitude, latitude, height,'
                         ' altitude, phone_app_latitude, phone_app_longitude, phone_app_x,'
                         ' phone_app_y, home_latitude,'
                         ' home_longitude, home_x, home_y, dt, frame_id, status'])
        for elem in model:
            writer.writerow(elem.values())

        return response

    @property
    def download_oldest_order(self):
        """download_oldest_order function"""
        model = SecondOrdersModel.objects.values().order_by('dt')
        response = HttpResponse(
            content_type='text/csv',
            headers={'Content-Disposition': 'attachment; filename="somefilename.csv"'},
        )

        writer = csv.writer(response)
        writer.writerow(['SELECT serial_no, product_type, longitude, latitude, height,'
                         ' altitude, phone_app_latitude, phone_app_longitude, phone_app_x,'
                         ' phone_app_y, home_latitude,'
                         ' home_longitude, home_x, home_y, dt, frame_id, status'])
        for elem in model:
            writer.writerow(elem.values())

        return response


class CombatLogic:
    """class for combat order page"""

    def __init__(self, date_search=None, fake_drone=None, get_time=None):
        self.date_search = date_search
        self.fake_drone = fake_drone
        self.get_time = get_time

        self.model_set = SecondOrdersModel.objects.filter(
            dt__icontains=self.date_search).values().exclude(
            serial_no=self.fake_drone).order_by('serial_no')

        if self.get_time is not None:
            self.model_set = SecondOrdersModel.objects.filter(
                dt__icontains=f"{self.date_search} {self.get_time[:2]}").values().exclude(
                serial_no=self.fake_drone
            ).order_by('serial_no')



    @property
    def today_req(self):
        """logic for today filter"""
        today_date = datetime.today().strftime('%y-%m-%d')

        model_set = SecondOrdersModel.objects.filter(
            dt__icontains=f"{today_date}").values().exclude(
            serial_no=self.fake_drone).order_by(
            'serial_no')

        if self.get_time is not None:
            model_set = SecondOrdersModel.objects.filter(
                dt__icontains=f"{today_date} {self.get_time[:3]}") \
                .values().exclude(
                serial_no=self.fake_drone).order_by(
                'serial_no')

        model = []
        total_quan = set()

        if len(model_set) == 1:
            quantity = SecondOrdersModel.objects.filter(dt__icontains=datetime.today()
                                                        .strftime('%y-%m-%d'),
                                                        serial_no=model_set[0][
                                                            'serial_no']).values().count()

            model_data = {'serial_no': model_set[0]['serial_no'],
                          'dt': model_set[0]['dt'].strftime('%m %d, %Y, %H:%M'),
                          'product_type': model_set[0]['product_type'],
                          'quantity': quantity,
                          'status': model_set[0]['status']
                          }
            total_quan.add(quantity)

            model.append(model_data)

        else:
            for num_elem, elem in enumerate(model_set):
                if num_elem == 0:
                    quantity = SecondOrdersModel.objects.filter(dt__icontains=datetime.today()
                                                                .strftime('%y-%m-%d'),
                                                                serial_no=model_set[num_elem][
                                                                    'serial_no']).values().count()

                    model_data = {'serial_no': model_set[num_elem]['serial_no'],
                                  'dt': model_set[num_elem]['dt'].strftime('%m %d, %Y, %H:%M'),
                                  'product_type': model_set[num_elem]['product_type'],
                                  'quantity': quantity,
                                  'status': model_set[num_elem]['status']
                                  }

                    model.append(model_data)
                    total_quan.add(quantity)
                else:
                    if model_set[num_elem]['serial_no'] != \
                            model_set[num_elem - 1]['serial_no']:
                        quantity = SecondOrdersModel.objects.filter(dt__icontains=datetime.today()
                                                                    .strftime('%y-%m-%d'),
                                                                    serial_no=model_set[num_elem][
                                                                        'serial_no']).values().count()

                        model_data = {'serial_no': model_set[num_elem]['serial_no'],
                                      'dt': model_set[num_elem]['dt'].strftime('%m %d, %Y, %H:%M'),
                                      'product_type': model_set[num_elem]['product_type'],
                                      'quantity': quantity,
                                      'status': model_set[num_elem]['status']
                                      }
                        model.append(model_data)
                        total_quan.add(quantity)

        # determination last 5 detection
        last_data_set_list = set()
        try:
            last_data_set = SecondOrdersModel.objects.values('serial_no').order_by('-id')[:5]
            for elem in last_data_set:
                if elem['serial_no'] not in last_data_set_list:
                    last_data_set_list.add(elem['serial_no'])
        except Exception:
            pass
        # ==============================================

        return [model, total_quan, last_data_set_list]

    @property
    def search_by_date(self):
        """logic for loading by date filter"""

        model = []

        if len(self.model_set) == 1:
            model.append(self.model_set)

        else:
            for elem in range(len(self.model_set)):
                if elem == 0:
                    quantity = SecondOrdersModel.objects.filter(dt__icontains=self.date_search,
                                                                serial_no=self.model_set[elem][
                                                                    'serial_no']).values().count()

                    model.append({'serial_no': self.model_set[elem]['serial_no'],
                                  'dt': self.model_set[elem]['dt'].strftime('%m %d, %Y, %H:%M'),
                                  'longitude': str(Decimal(self.model_set[elem]['longitude'])),
                                  'latitude': str(Decimal(self.model_set[elem]['latitude'])),
                                  'product_type': self.model_set[elem]['product_type'],
                                  'quantity': quantity,
                                  'status': self.model_set[elem]['status']
                                  })

                else:
                    if self.model_set[elem]['serial_no'] != self.model_set[elem - 1]['serial_no']:
                        quantity = SecondOrdersModel.objects.filter(dt__icontains=self.date_search,
                                                                    serial_no=self.model_set[elem][
                                                                        'serial_no']).values().count()

                        model.append({'serial_no': self.model_set[elem]['serial_no'],
                                      'dt': self.model_set[elem]['dt'].strftime('%m %d, %Y, %H:%M'),
                                      'longitude': str(Decimal(self.model_set[elem]['longitude'])),
                                      'latitude': str(Decimal(self.model_set[elem]['latitude'])),
                                      'product_type': self.model_set[elem]['product_type'],
                                      'quantity': quantity,
                                      'status': self.model_set[elem]['status']
                                      })
            return model

    @property
    def search_by_today_statistics(self):
        """built data for today statistics order """
        cur_date = datetime.today().strftime("%y-%m-%d")
        cur_month = datetime.today().strftime("%y-%m")
        tommorow_date = (datetime.today() - timedelta(days=1)).strftime("%d")

        model_set = SecondOrdersModel.objects.filter(
            Q(dt__icontains=cur_date) | Q(dt__icontains=f"{cur_month}-{tommorow_date} 15:") | Q(
                dt__icontains=f"{cur_month}-{tommorow_date} 16:") | Q(
                dt__icontains=f"{cur_month}-{tommorow_date} 17:") | Q(
                dt__icontains=f"{cur_month}-{tommorow_date} 18:") | Q(
                dt__icontains=f"{cur_month}-{tommorow_date} 19:") | Q(
                dt__icontains=f"{cur_month}-{tommorow_date} 20:") | Q(
                dt__icontains=f"{cur_month}-{tommorow_date} 21:") | Q(
                dt__icontains=f"{cur_month}-{tommorow_date} 22:") | Q(
                dt__icontains=f"{cur_month}-{tommorow_date} 23:")).values().exclude(
            serial_no=self.fake_drone
        ).order_by('serial_no')

        model = []

        if len(model_set) == 1:
            model.append(self.model_set)

        else:
            for iter_num, value in enumerate(model_set):
                if iter_num == 0:
                    quantity = SecondOrdersModel.objects.filter(
                        Q(dt__icontains=cur_date) | Q(dt__icontains=f"{cur_month}-{tommorow_date} 15:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 16:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 17:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 18:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 19:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 20:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 21:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 22:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 23:"),
                        serial_no=model_set[iter_num][
                            'serial_no']).values().count()

                    model.append({'serial_no': model_set[iter_num]['serial_no'],
                                  'dt': model_set[iter_num]['dt'].strftime('%m %d, %Y, %H:%M'),
                                  'longitude': str(Decimal(model_set[iter_num]['longitude'])),
                                  'latitude': str(Decimal(model_set[iter_num]['latitude'])),
                                  'product_type': model_set[iter_num]['product_type'],
                                  'quantity': quantity,
                                  'status': model_set[iter_num]['status']
                                  })

                else:
                    if model_set[iter_num]['serial_no'] != model_set[iter_num - 1]['serial_no']:
                        quantity = SecondOrdersModel.objects.filter(
                            Q(dt__icontains=cur_date) | Q(dt__icontains=f"{cur_month}-{tommorow_date} 15:") | Q(
                                dt__icontains=f"{cur_month}-{tommorow_date} 16:") | Q(
                                dt__icontains=f"{cur_month}-{tommorow_date} 17:") | Q(
                                dt__icontains=f"{cur_month}-{tommorow_date} 18:") | Q(
                                dt__icontains=f"{cur_month}-{tommorow_date} 19:") | Q(
                                dt__icontains=f"{cur_month}-{tommorow_date} 20:") | Q(
                                dt__icontains=f"{cur_month}-{tommorow_date} 21:") | Q(
                                dt__icontains=f"{cur_month}-{tommorow_date} 22:") | Q(
                                dt__icontains=f"{cur_month}-{tommorow_date} 23:"),
                            serial_no=model_set[iter_num][
                                'serial_no']).values().count()

                        model.append({'serial_no': model_set[iter_num]['serial_no'],
                                      'dt': model_set[iter_num]['dt'].strftime('%m %d, %Y, %H:%M'),
                                      'longitude': str(Decimal(model_set[iter_num]['longitude'])),
                                      'latitude': str(Decimal(model_set[iter_num]['latitude'])),
                                      'product_type': model_set[iter_num]['product_type'],
                                      'quantity': quantity,
                                      'status': model_set[iter_num]['status']
                                      })

        return model

    @property
    def search_by_week_statistics(self):
        day_1 = (datetime.today() - timedelta(days=1)).strftime("%y-%m-%d")
        day_2 = (datetime.today() - timedelta(days=2)).strftime("%y-%m-%d")
        day_3 = (datetime.today() - timedelta(days=3)).strftime("%y-%m-%d")
        day_4 = (datetime.today() - timedelta(days=4)).strftime("%y-%m-%d")
        day_5 = (datetime.today() - timedelta(days=5)).strftime("%y-%m-%d")
        day_6 = (datetime.today() - timedelta(days=6)).strftime("%y-%m-%d")
        day_7 = (datetime.today() - timedelta(days=7)).strftime("%y-%m-%d")

        model_set = SecondOrdersModel.objects.filter(
            Q(dt__icontains=day_1) | Q(dt__icontains=day_2) | Q(
                dt__icontains=day_3) | Q(
                dt__icontains=day_4) | Q(
                dt__icontains=day_5) | Q(
                dt__icontains=day_6) | Q(
                dt__icontains=day_7)).values().exclude(
            serial_no=self.fake_drone
        ).order_by('serial_no')

        model = []

        if len(model_set) == 1:
            model.append(self.model_set)

        else:
            for num_iter, values in enumerate(model_set):
                if num_iter == 0:
                    quantity = SecondOrdersModel.objects.filter(
                        Q(dt__icontains=day_1) | Q(dt__icontains=day_2) | Q(
                            dt__icontains=day_3) | Q(
                            dt__icontains=day_4) | Q(
                            dt__icontains=day_5) | Q(
                            dt__icontains=day_6) | Q(
                            dt__icontains=day_7),
                        serial_no=model_set[num_iter][
                            'serial_no']).values().count()

                    model.append({'serial_no': model_set[num_iter]['serial_no'],
                                  'dt': model_set[num_iter]['dt'].strftime('%m %d, %Y, %H:%M'),
                                  'longitude': str(Decimal(model_set[num_iter]['longitude'])),
                                  'latitude': str(Decimal(model_set[num_iter]['latitude'])),
                                  'product_type': model_set[num_iter]['product_type'],
                                  'quantity': quantity,
                                  'status': model_set[num_iter]['status']
                                  })

                else:
                    if model_set[num_iter]['serial_no'] != model_set[num_iter - 1]['serial_no']:
                        quantity = SecondOrdersModel.objects.filter(
                            Q(dt__icontains=day_1) | Q(dt__icontains=day_2) | Q(
                                dt__icontains=day_3) | Q(
                                dt__icontains=day_4) | Q(
                                dt__icontains=day_5) | Q(
                                dt__icontains=day_6) | Q(
                                dt__icontains=day_7),
                            serial_no=model_set[num_iter][
                                'serial_no']).values().count()

                        model.append({'serial_no': model_set[num_iter]['serial_no'],
                                      'dt': model_set[num_iter]['dt'].strftime('%m %d, %Y, %H:%M'),
                                      'longitude': str(Decimal(model_set[num_iter]['longitude'])),
                                      'latitude': str(Decimal(model_set[num_iter]['latitude'])),
                                      'product_type': model_set[num_iter]['product_type'],
                                      'quantity': quantity,
                                      'status': model_set[num_iter]['status']
                                      })

        return model


# change status drone on combat_data page

class ChoseStatusCombat:

    def __init__(self, status=None):
        self.status = status

    def change_status(self):
        year = self.status[4].split(',')[0]
        month = self.status[2].split(',')[0]
        day = self.status[3].split(',')[0]
        whom = self.status[1]
        SecondOrdersModel.objects.filter(serial_no=self.status[0],
                                         dt__icontains=f"{year}-{month}-{day}").update(
            status=whom)


class OpenDataCombatLogicClass:
    """OpenData get request in combat logic"""

    def __init__(self, *args):
        self.input_data = args[0]

        self.serial_no = self.input_data.split(' ')[0]
        self.current_year = self.input_data.split(',')[1].split(' ')[1]
        self.current_month = self.input_data.split(' ')[1]
        self.current_day = self.input_data.split(' ')[2].split(',')[0]
        self.date_set = {
            'January': '01',
            'February': '02',
            'March': '03',
            'April': '04',
            'May': '05',
            'June': '06',
            'July': '07',
            'August': '08',
            'September': '09',
            'October': '10',
            'November': '11',
            'December': '12'
        }

    @property
    def enter_to_detail_data(self):
        """entering to detail data in combar order page """
        if int(self.current_day) < 10:
            model_detail = SecondOrdersModel.objects.filter(
                dt__icontains=f"{self.current_year}-{self.current_month}-{self.current_day}",
                serial_no=self.serial_no).values().order_by(
                '-dt')
            model = SecondOrdersModel.objects.filter(
                dt__icontains=f"{self.current_year}-{self.current_month}-{self.current_day}",
                serial_no=self.serial_no).values().order_by(
                '-dt')[0]

            data = {
                'model': model,
                'product_type': model['product_type'],
                'model_detail': model_detail,
                'action': 1,

            }

            return data

        else:
            model_detail = SecondOrdersModel.objects.filter(
                dt__icontains=f"{self.current_year}-{self.current_month}-{self.current_day}",
                serial_no=self.serial_no).values().order_by(
                '-dt')
            model = SecondOrdersModel.objects.filter(
                dt__icontains=f"{self.current_year}-{self.current_month}-{self.current_day}",
                serial_no=self.serial_no).values().order_by(
                '-dt')[0]

            data = {
                'model': model,
                'model_detail': model_detail,
                'action': 1,

            }

            return data


class BuildCombatOrders:
    """Combat orders download docx logic"""

    @property
    def build_orders(self):
        """logic for building order docx on CombatOrders page"""

        logic = BuildStatistics(datetime.today().strftime("%y-%m-%d")).today_statistics_order

        drones = {}
        drone_str = ''

        for key, value in logic.items():
            if value != 0:
                if key == 'total_value':
                    pass
                elif key == 'dirty_total_value':
                    pass
                elif key == 'ally':
                    drones['свої'] = value
                    drone_str += f"\t * свої - {value}\n"
                elif key == 'fag':
                    drones['ворожі'] = value
                    drone_str += f"\t * ворожі - {value}\n"
                elif key == 'unknown':
                    drones['fake_gps'] = value
                    drone_str += f"\t * fake_gps - {value}\n"
                else:
                    drones[key] = value
                    drone_str += f"\t * {key} - {value}\n"

        if len(logic) == 0:
            text = Document()

            text.add_paragraph("""Немає данних за обраний період""")

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': 'attachment; filename="combat_orders.docx"'},
            )
            text.save(response)
            return response

        else:
            text = Document()
            text.add_paragraph(f"""На південно-східних околицях н.п. Дронівка встановлено систему Джміль для виявлення БпЛА виробництва DJI:
• За звітній період було виявлено {logic['total_value']} унікальних БпЛА, кількість детекцій за добу {logic['dirty_total_value']};
• Було виявлено наступний модельний ряд та кількість:
    {drone_str}
""")

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': 'attachment; filename="combat_orders.docx"'},
            )
            text.save(response)
            return response


class MainPageLogic:
    """logic for main page"""

    def __init__(self):
        pass

    @property
    def total_month_result(self):
        """count total month result"""
        total_value = SecondOrdersModel.objects.filter(
            dt__icontains=datetime.today().strftime('%y-%m')).aggregate(
            Count('serial_no'))

        return total_value


class BuildStatistics:

    def __init__(self, current_day):
        self.logic = CombatLogic(datetime.today().strftime('%y-%m'))
        self.data_set = {}
        self.model = []
        self.open_data = current_day

    @property
    def total_results_for_chosen_month(self):
        data_set = CombatLogic(self.open_data).search_by_date

        data_1 = {
            41: 0,
            44: 0,
            53: 0,
            58: 0,
            60: 0,
            63: 0,
            66: 0,
            67: 0,
            68: 0,
            69: 0,
            70: 0,
            73: 0,
            77: 0,
            86: 0,
            'ally': 0,
            'fag': 0,
            'unknown': 0
        }

        for el in data_set:
            if el['product_type'] in data_1.keys():
                data_1[el['product_type']] += 1
                if any([el['latitude'].startswith('48.9731'),
                        el['latitude'].startswith('48.9732'),
                        el['latitude'].startswith('48.9733'),
                        el['latitude'].startswith('48.9734'),
                        el['latitude'].startswith('48.9735'),
                        el['latitude'].startswith('48.961'),
                        el['latitude'].startswith('48.962'),
                        el['latitude'].startswith('48.963'),
                        el['latitude'].startswith('48.964'),
                        el['latitude'].startswith('48.965'),
                        el['latitude'].startswith('48.966'),
                        el['latitude'].startswith('48.956'),
                        el['latitude'].startswith('48.955'),
                        el['latitude'].startswith('48.954'),
                        el['latitude'].startswith('48.953'),
                        el['latitude'].startswith('48.952'),
                        el['latitude'].startswith('48.951'),
                        ]):
                    data_1['ally'] += 1
                elif el['longitude'].startswith('-'):
                    data_1['unknown'] += 1
                else:
                    data_1['fag'] += 1
            else:
                data_1[68] += 1

        data = {'total_value': len(data_set),
                'dirty_total_value': len(
                    SecondOrdersModel.objects.filter(dt__icontains=self.open_data)),
                'mavic_2': data_1[41],
                'M_200_v2': data_1[44],
                'Mavic_Mini': data_1[53],
                'Mavic_Air_2': data_1[58],
                'M300RTK': data_1[60],
                'mini_2': data_1[63],
                'air_2s': data_1[66],
                'm30': data_1[67],
                'mavic_3': data_1[68],
                'mavic2Enterprise': data_1[69],
                'mini_se': data_1[70],
                'mini_3_Pro': data_1[73],
                'Mavic_3T_3E': data_1[77],
                'Mavic_3_Classic': data_1[86],
                'ally': data_1['ally'],
                'fag': data_1['fag'],
                'unknown': data_1['unknown'],
                }
        return data

    @property
    def today_statistics_order(self):
        """built statistics by today in statistics page """
        data_set = CombatLogic(self.open_data).search_by_today_statistics
        cur_date = datetime.today().strftime("%y-%m-%d")
        cur_month = datetime.today().strftime("%y-%m")
        tommorow_date = (datetime.today() - timedelta(days=1)).strftime("%d")

        data_1 = {
            0: 0,
            41: 0,
            44: 0,
            53: 0,
            58: 0,
            60: 0,
            63: 0,
            66: 0,
            67: 0,
            68: 0,
            69: 0,
            70: 0,
            73: 0,
            77: 0,
            86: 0,
            'ally': 0,
            'fag': 0,
            'unknown': 0
        }

        for el in data_set:
            if el['product_type'] in data_1.keys():
                data_1[el['product_type']] += 1
                if any([el['latitude'].startswith('48.9731'),
                        el['latitude'].startswith('48.9732'),
                        el['latitude'].startswith('48.9733'),
                        el['latitude'].startswith('48.9734'),
                        el['latitude'].startswith('48.9735'),
                        el['latitude'].startswith('48.961'),
                        el['latitude'].startswith('48.962'),
                        el['latitude'].startswith('48.963'),
                        el['latitude'].startswith('48.964'),
                        el['latitude'].startswith('48.965'),
                        el['latitude'].startswith('48.966'),
                        el['latitude'].startswith('48.956'),
                        el['latitude'].startswith('48.955'),
                        el['latitude'].startswith('48.954'),
                        el['latitude'].startswith('48.953'),
                        el['latitude'].startswith('48.952'),
                        el['latitude'].startswith('48.951'),
                        ]):
                    data_1['ally'] += 1
                elif el['longitude'].startswith('-'):
                    data_1['unknown'] += 1
                else:
                    data_1['fag'] += 1
            else:
                data_1[68] += 1

        data = {'total_value': len(data_set),
                'dirty_total_value': len(
                    SecondOrdersModel.objects.filter(
                        Q(dt__icontains=cur_date) |
                        Q(dt__icontains=f"{cur_month}-{tommorow_date} 15:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 16:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 17:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 18:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 19:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 20:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 21:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 22:") | Q(
                            dt__icontains=f"{cur_month}-{tommorow_date} 23:"))),
                '1001_frimware': data_1[0],
                'mavic_2': data_1[41],
                'M_200_v2': data_1[44],
                'Mavic_Mini': data_1[53],
                'Mavic_Air_2': data_1[58],
                'M300RTK': data_1[60],
                'mini_2': data_1[63],
                'air_2s': data_1[66],
                'm30': data_1[67],
                'mavic_3': data_1[68],
                'mavic2Enterprise': data_1[69],
                'mini_se': data_1[70],
                'mini_3_Pro': data_1[73],
                'Mavic_3T_3E': data_1[77],
                'Mavic_3_Classic': data_1[86],
                'ally': data_1['ally'],
                'fag': data_1['fag'],
                'unknown': data_1['unknown'],
                }
        return data

    @property
    def weak_statistics_order(self):
        """built statistics by the week on statistics page"""
        data_set = CombatLogic(self.open_data).search_by_week_statistics
        day_1 = (datetime.today() - timedelta(days=1)).strftime("%y-%m-%d")
        day_2 = (datetime.today() - timedelta(days=2)).strftime("%y-%m-%d")
        day_3 = (datetime.today() - timedelta(days=3)).strftime("%y-%m-%d")
        day_4 = (datetime.today() - timedelta(days=4)).strftime("%y-%m-%d")
        day_5 = (datetime.today() - timedelta(days=5)).strftime("%y-%m-%d")
        day_6 = (datetime.today() - timedelta(days=6)).strftime("%y-%m-%d")
        day_7 = (datetime.today() - timedelta(days=7)).strftime("%y-%m-%d")

        date_data_set = {'day_1': day_1,
                         'day_2': day_2,
                         'day_3': day_3,
                         'day_4': day_4,
                         'day_5': day_5,
                         'day_6': day_6,
                         'day_7': day_7,
                         }
        quan_by_date = {'day_1': len(SecondOrdersModel.objects.filter(dt__icontains=day_1)),
                        'day_2': len(SecondOrdersModel.objects.filter(dt__icontains=day_2)),
                        'day_3': len(SecondOrdersModel.objects.filter(dt__icontains=day_3)),
                        'day_4': len(SecondOrdersModel.objects.filter(dt__icontains=day_4)),
                        'day_5': len(SecondOrdersModel.objects.filter(dt__icontains=day_5)),
                        'day_6': len(SecondOrdersModel.objects.filter(dt__icontains=day_6)),
                        'day_7': len(SecondOrdersModel.objects.filter(dt__icontains=day_7))
                        }
        data_1 = {
            0: 0,
            41: 0,
            44: 0,
            53: 0,
            58: 0,
            60: 0,
            63: 0,
            66: 0,
            67: 0,
            68: 0,
            69: 0,
            70: 0,
            73: 0,
            77: 0,
            86: 0,
            'ally': 0,
            'fag': 0,
            'unknown': 0
        }

        for el in data_set:
            if el['product_type'] in data_1.keys():
                data_1[el['product_type']] += 1
                if any([el['latitude'].startswith('48.9731'),
                        el['latitude'].startswith('48.9732'),
                        el['latitude'].startswith('48.9733'),
                        el['latitude'].startswith('48.9734'),
                        el['latitude'].startswith('48.9735'),
                        el['latitude'].startswith('48.961'),
                        el['latitude'].startswith('48.962'),
                        el['latitude'].startswith('48.963'),
                        el['latitude'].startswith('48.964'),
                        el['latitude'].startswith('48.965'),
                        el['latitude'].startswith('48.966'),
                        el['latitude'].startswith('48.956'),
                        el['latitude'].startswith('48.955'),
                        el['latitude'].startswith('48.954'),
                        el['latitude'].startswith('48.953'),
                        el['latitude'].startswith('48.952'),
                        el['latitude'].startswith('48.951'),
                        ]):
                    data_1['ally'] += 1
                elif el['longitude'].startswith('-'):
                    data_1['unknown'] += 1
                else:
                    data_1['fag'] += 1
            else:
                data_1[68] += 1

        data = {'total_value': len(data_set),
                'dirty_total_value': len(
                    SecondOrdersModel.objects.filter(
                        Q(dt__icontains=day_1) | Q(dt__icontains=day_2) | Q(
                            dt__icontains=day_3) | Q(
                            dt__icontains=day_4) | Q(
                            dt__icontains=day_5) | Q(
                            dt__icontains=day_6) | Q(
                            dt__icontains=day_7))),
                '1001_firmware': data_1[0],
                'mavic_2': data_1[41],
                'M_200_v2': data_1[44],
                'Mavic_Mini': data_1[53],
                'Mavic_Air_2': data_1[58],
                'M300RTK': data_1[60],
                'mini_2': data_1[63],
                'air_2s': data_1[66],
                'm30': data_1[67],
                'mavic_3': data_1[68],
                'mavic2Enterprise': data_1[69],
                'mini_se': data_1[70],
                'mini_3_Pro': data_1[73],
                'Mavic_3T_3E': data_1[77],
                'Mavic_3_Classic': data_1[86],
                'ally': data_1['ally'],
                'fag': data_1['fag'],
                'unknown': data_1['unknown'],
                }
        return [data, date_data_set, quan_by_date]

    # function for rating onn main page

    @property
    def top_rank(self):
        """built top rank order in main page"""
        logic = self.logic.today_req

        if len(logic[1]) == 0:
            return 'No data'
        elif len(logic[1]) > 6:
            top = sorted(logic[1])[-6:]
        else:
            top = sorted(logic[1])

        for el in logic[0]:
            if el['quantity'] in top:
                self.data_set = {'serial_no': el['serial_no'],
                                 'quantity': el['quantity'],
                                 'dt': el['dt']
                                 }
                self.model.append(self.data_set)
        return self.model


class LogicAnalyze:

    def __init__(self, data_1, data_2):
        self.date_1 = data_1
        self.date_2 = data_2
        self.data_set_1 = BuildStatistics(self.date_1[:7])
        self.data_set_2 = BuildStatistics(self.date_2[:7])
        # self.data = {}
        self.data = {
            'dirty_total_value': 0,
            41: 0,
            44: 0,
            53: 0,
            58: 0,
            60: 0,
            63: 0,
            66: 0,
            67: 0,
            68: 0,
            69: 0,
            70: 0,
            73: 0,
            77: 0,
            86: 0,
        }

    @property
    def make_anylyze(self):

        # if all([(self.data_set_1.total_results_for_chosen_month['total_value'] != 0),
        #        (self.data_set_1.total_results_for_chosen_month['total_value'] != 0)]):
        #    total_value = (self.data_set_1.total_results_for_chosen_month['total_value'] /
        #                   self.data_set_2.total_results_for_chosen_month['total_value']) * 100
        #    self.data['total_value'] = f"{round(total_value, 2)} %"

        if all([(len(SecondOrdersModel.objects.filter(dt__icontains=self.date_1)) != 0),
                (len(SecondOrdersModel.objects.filter(dt__icontains=self.date_2)) != 0)]):
            dirty_total_value = (len(SecondOrdersModel.objects.filter(dt__icontains=self.date_1)) /
                                 (len(SecondOrdersModel.objects.filter(dt__icontains=self.date_1)))) * 100
            self.data['dirty_total_value'] += round(dirty_total_value, 2)

        for el in self.data:
            if el != 'dirty_total_value':
                if all([(len(SecondOrdersModel.objects.filter(dt__icontains=self.date_1, product_type=el)) != 0),
                        (len(SecondOrdersModel.objects.filter(dt__icontains=self.date_2, product_type=el)) != 0)]):
                    self.data[el] += ((
                            len(SecondOrdersModel.objects.filter(dt__icontains=self.date_1, product_type=el)) /
                            (len(SecondOrdersModel.objects.filter(dt__icontains=self.date_2,
                                                                  product_type=el))))) * 100

                elif len(SecondOrdersModel.objects.filter(dt__icontains=self.date_1, product_type=el)) == 0:
                    self.data[el] += len(SecondOrdersModel.objects.filter(dt__icontains=self.date_2, product_type=el))

                elif len(SecondOrdersModel.objects.filter(dt__icontains=self.date_2, product_type=el)) == 0:
                    self.data[el] += len(SecondOrdersModel.objects.filter(dt__icontains=self.date_1, product_type=el))

        final_data = {
            'dirty_total_value': self.data['dirty_total_value'],
            'mavic_2': round(self.data[41], 2),
            'M_200_v2': round(self.data[44], 2),
            'Mavic_Mini': round(self.data[53], 2),
            'Mavic_Air_2': round(self.data[58], 2),
            'M300RTK': round(self.data[60], 2),
            'mini_2': round(self.data[63], 2),
            'air_2s': round(self.data[66], 2),
            'm30': round(self.data[67], 2),
            'mavic_3': round(self.data[68], 2),
            'mavic2Enterprise': round(self.data[69], 2),
            'mini_se': round(self.data[70], 2),
            'mini_3_Pro': round(self.data[73], 2),
            'Mavic_3T_3E': round(self.data[77], 2),
            'Mavic_3_Classic': round(self.data[86], 2),

        }

        return self.data


class AddFlightRecorderData:

    def __init__(self, drone_type, drone_id, coord_x, coord_y):
        self.drona_type = drone_type
        self.drone_id = drone_id
        self.coord_x = coord_x
        self.coord_y = coord_y

    def add_data(self):
        FlightRecorderModel.objects.create(drone_type=self.drona_type, drone_id=self.drone_id, coord_x=self.coord_x,
                                           coord_y=self.coord_y)


class FilterFlightRecordData:

    def __init__(self, date_search=None, find_time=None, today=None):
        self.date_search = date_search
        self.find_time = find_time
        self.today = today

    def find_by_today_filter(self):
        data = FlightRecorderModel.objects.filter(record_data__icontains=self.today)
        return data


class SkySafeLogic:

    def __init__(self, date_search=None, fake_drone=None, get_time=None):
        self.date_search = date_search
        self.fake_drone = fake_drone
        self.get_time = get_time

        if self.get_time is not None:
            self.model_set = SkySafeData.objects.filter(
                write_time__icontains=f"{self.date_search} {self.get_time[:3]}").values().exclude(
                persistent_id=self.fake_drone
            ).order_by('persistent_id')
        else:
            self.model_set = SkySafeData.objects.filter(
                write_time__icontains=self.date_search).values().exclude(
                persistent_id=self.fake_drone).order_by('persistent_id')

    @property
    def today_req(self):

        model = []
        total_quan = []

        if len(self.model_set) == 1:
            quantity = SkySafeData.objects.filter(write_time__icontains=datetime.today().strftime('%y-%m-%d'),
                                                  persistent_id=self.model_set[0][
                                                      'persistent_id']).values().count()

            model_data = {'serial_no': self.model_set[0]['persistent_id'],
                          'dt': self.model_set[0]['write_time'].strftime('%m %d, %Y, %H:%M'),
                          'product_type': self.model_set[0]['tgt_model'],
                          'quantity': quantity,
                          }
            total_quan.append(quantity)

            model.append(model_data)

        else:
            for el in range(len(self.model_set)):
                if el == 0:
                    quantity = SkySafeData.objects.filter(write_time__icontains=datetime.today().strftime('%y-%m-%d'),
                                                          persistent_id=self.model_set[el][
                                                              'persistent_id']).values().count()

                    model_data = {'serial_no': self.model_set[0]['persistent_id'],
                                  'dt': self.model_set[0]['write_time'].strftime('%m %d, %Y, %H:%M'),
                                  'product_type': self.model_set[0]['tgt_model'],
                                  'quantity': quantity,

                                  }

                    model.append(model_data)
                    total_quan.append(quantity)
                else:
                    if self.model_set[el]['persistent_id'] != self.model_set[el - 1]['persistent_id']:
                        quantity = SkySafeData.objects.filter(
                            write_time__icontains=datetime.today().strftime('%y-%m-%d'),
                            persistent_id=self.model_set[el][
                                'persistent_id']).values().count()

                        model_data = {'serial_no': self.model_set[el]['persistent_id'],
                                      'dt': self.model_set[el]['write_time'].strftime('%m %d, %Y, %H:%M'),
                                      'product_type': self.model_set[el]['tgt_model'],
                                      'quantity': quantity,
                                      }

                        model.append(model_data)
                        total_quan.append(quantity)
        return [model, total_quan, self.model_set, self.date_search, self.get_time]

    @property
    def search_by_date(self):

        model = []

        if len(self.model_set) == 1:
            model.append(self.model_set)

        else:
            for el in range(len(self.model_set)):
                if el == 0:
                    quantity = SkySafeData.objects.filter(write_time__icontains=self.date_search,
                                                          persistent_id=self.model_set[el][
                                                              'persistent_id']).values().count()

                    model.append({'serial_no': self.model_set[el]['persistent_id'],
                                  'dt': self.model_set[el]['write_time'].strftime('%m %d, %Y, %H:%M'),
                                  'product_type': self.model_set[el]['tgt_model'],
                                  'quantity': quantity,

                                  })

                else:
                    if self.model_set[el]['persistent_id'] != self.model_set[el - 1]['persistent_id']:
                        quantity = SkySafeData.objects.filter(write_time__icontains=self.date_search,
                                                              persistent_id=self.model_set[el][
                                                                  'persistent_id']).values().count()

                        model.append({'serial_no': self.model_set[el]['persistent_id'],
                                      'dt': self.model_set[el]['write_time'].strftime('%m %d, %Y, %H:%M'),
                                      'product_type': self.model_set[el]['tgt_model'],
                                      'quantity': quantity,

                                      })

            return model


'''OpenData get request in SkySafe logic'''


class OpenDataSkySafeClass:
    def __init__(self, *args):
        self.input_data = args[0]

        self.serial_no = self.input_data.split(' ')[0]
        self.current_year = self.input_data.split(',')[1].split(' ')[1]
        self.current_month = self.input_data.split(' ')[1]
        self.current_day = self.input_data.split(' ')[2].split(',')[0]
        self.date_set = {
            'January': '01',
            'February': '02',
            'March': '03',
            'April': '04',
            'May': '05',
            'June': '06',
            'July': '07',
            'August': '08',
            'September': '09',
            'October': '10',
            'November': '11',
            'December': '12'
        }

    @property
    def enter_to_detail_data(self):
        if int(self.current_day) < 10:
            model_detail = SkySafeData.objects.filter(
                write_time__icontains=f"{self.current_year}-{self.current_month}-{self.current_day}",
                persistent_id=self.serial_no).values().order_by(
                '-write_time')
            model = SkySafeData.objects.filter(
                write_time__icontains=f"{self.current_year}-{self.current_month}-{self.current_day}",
                persistent_id=self.serial_no).values().order_by(
                '-write_time')[0]

            data = {
                'model': model,
                'product_type': model['persistent_id'],
                'model_detail': model_detail,
                'action': 1,

            }

            return data

        else:
            model_detail = SkySafeData.objects.filter(
                write_time__icontains=f"{self.current_year}-{self.current_month}-{self.current_day}",
                persistent_id=self.serial_no).values().order_by(
                '-write_time')
            model = SkySafeData.objects.filter(
                write_time__icontains=f"{self.current_year}-{self.current_month}-{self.current_day}",
                persistent_id=self.serial_no).values().order_by(
                '-write_time')[0]

            data = {
                'model': model,
                'model_detail': model_detail,
                'action': 1,

            }

            return data
